import json
import pypdf
import docx
import os
from django.utils import timezone
from django.core.signing import TimestampSigner, SignatureExpired, BadSignature
from django.core.mail import send_mail
from django.http import FileResponse, HttpResponse
from django.db.models import Q
from django.conf import settings
from rest_framework import status, permissions, generics
from rest_framework.views import APIView
from rest_framework.response import Response

from .models import Order, PricingConfig
from .serializers import OrderSerializer, PricingConfigSerializer
from .pricing import default_pricing_kwargs, gst_breakdown, package_catalog, PACKAGE_LABELS, resolve_package
from accounts.models import User
from accounts.permissions import IsSuperAdmin, IsCollegeAdmin
from colleges.models import College

# Create signer for secure download link
signer = TimestampSigner()
MAX_UPLOAD_BYTES = 25 * 1024 * 1024
ALLOWED_EXTENSIONS = ('.pdf', '.doc', '.docx')


def extract_submission_meta(request):
    raw_co_authors = request.data.get('co_authors') or '[]'
    if isinstance(raw_co_authors, str):
        try:
            co_authors = json.loads(raw_co_authors)
        except json.JSONDecodeError:
            co_authors = []
    elif isinstance(raw_co_authors, list):
        co_authors = raw_co_authors
    else:
        co_authors = []

    cleaned = []
    for item in co_authors:
        if not isinstance(item, dict):
            continue
        name = str(item.get('name') or '').strip()
        if not name:
            continue
        cleaned.append({
            'name': name,
            'email': str(item.get('email') or '').strip(),
            'institution': str(item.get('institution') or '').strip(),
        })

    return {
        'paper_title': str(request.data.get('paper_title') or '').strip(),
        'paper_type': str(request.data.get('paper_type') or '').strip(),
        'subject_area': str(request.data.get('subject_area') or '').strip(),
        'purpose': str(request.data.get('purpose') or '').strip(),
        'keywords': str(request.data.get('keywords') or '').strip(),
        'author_name': str(request.data.get('author_name') or '').strip(),
        'author_email': str(request.data.get('author_email') or '').strip(),
        'author_institution': str(request.data.get('author_institution') or '').strip(),
        'author_country': str(request.data.get('author_country') or '').strip(),
        'co_authors': cleaned,
    }

def get_word_count(file):
    name = file.name.lower()
    text = ""
    try:
        # Seek file to start
        file.seek(0)
        if name.endswith('.pdf'):
            reader = pypdf.PdfReader(file)
            for page in reader.pages:
                text += page.extract_text() or ""
        elif name.endswith('.docx'):
            doc = docx.Document(file)
            for para in doc.paragraphs:
                text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
        else:
            text = file.read().decode('utf-8', errors='ignore')
    except Exception as e:
        print(f"Error parsing file for word count: {e}")
        text = "fallback text"
    
    words = text.split()
    count = len(words)
    return max(count, 1)


class WordCountEstimateView(APIView):
    permission_classes = [permissions.AllowAny]

    def post(self, request):
        if 'file' not in request.FILES:
            return Response({"error": "No file uploaded"}, status=status.HTTP_400_BAD_REQUEST)
        
        uploaded_file = request.FILES['file']
        word_count = get_word_count(uploaded_file)

        config = PricingConfig.objects.last()
        if not config:
            config = PricingConfig.objects.create(**default_pricing_kwargs())

        package, total_price, _flags = resolve_package(request.data, config)
        taxable, gst, gross = gst_breakdown(total_price)

        return Response({
            "filename": uploaded_file.name,
            "word_count": word_count,
            "package": package,
            "package_label": PACKAGE_LABELS[package],
            "taxable_value": float(taxable),
            "gst": float(gst),
            "total_price": float(gross),
        })


class OrderListCreateView(generics.ListCreateAPIView):
    serializer_class = OrderSerializer

    def get_queryset(self):
        user = self.request.user
        if user.role == 'super_admin':
            queryset = Order.objects.exclude(status='Pending Payment')
            search_query = self.request.query_params.get('search', '').strip()
            status_query = self.request.query_params.get('status', '').strip()
            if search_query:
                search_filters = (
                    Q(document__icontains=search_query)
                    | Q(paper_title__icontains=search_query)
                    | Q(author_name__icontains=search_query)
                    | Q(user__username__icontains=search_query)
                    | Q(user__email__icontains=search_query)
                    | Q(payment__razorpay_payment_id__icontains=search_query)
                    | Q(payment__razorpay_order_id__icontains=search_query)
                    | Q(payment__transaction_id__icontains=search_query)
                )
                if search_query.isdigit():
                    search_filters |= Q(user__id=int(search_query)) | Q(id=int(search_query))
                queryset = queryset.filter(search_filters)
            if status_query:
                queryset = queryset.filter(status__iexact=status_query)
            return queryset.select_related('user', 'college', 'payment').order_by('-created_at')
        elif user.role == 'college_admin':
            if not user.college:
                return Order.objects.none()
            queryset = Order.objects.filter(college=user.college)
            department = self.request.query_params.get('department')
            start_date = self.request.query_params.get('start_date')
            end_date = self.request.query_params.get('end_date')
            min_similarity = self.request.query_params.get('min_similarity')
            max_similarity = self.request.query_params.get('max_similarity')
            
            if department:
                queryset = queryset.filter(department__iexact=department)
            if start_date:
                queryset = queryset.filter(created_at__date__gte=start_date)
            if end_date:
                queryset = queryset.filter(created_at__date__lte=end_date)
            if min_similarity:
                queryset = queryset.filter(similarity_score__gte=float(min_similarity))
            if max_similarity:
                queryset = queryset.filter(similarity_score__lte=float(max_similarity))
                
            return queryset.select_related('user', 'college', 'payment').order_by('-created_at')
        else:
            return Order.objects.filter(user=user).select_related('payment', 'college').order_by('-created_at')

    def list(self, request, *args, **kwargs):
        queryset = self.get_queryset()
        serializer = self.get_serializer(queryset, many=True)
        data = serializer.data

        # Fallback to Firestore if local queryset is empty
        if not data and request.user.is_authenticated:
            try:
                from services.firestore_service import get_firestore_orders_for_user
                fs_orders = get_firestore_orders_for_user(request.user.id)
                if fs_orders:
                    return Response(fs_orders)
            except Exception as e:
                print(f"Firestore Read Error: {e}")

        return Response(data)

    def create(self, request, *args, **kwargs):
        user = request.user
        if 'document' not in request.FILES:
            return Response({"error": "Document file is required"}, status=status.HTTP_400_BAD_REQUEST)
            
        file = request.FILES['document']
        filename = (file.name or '').lower()
        if not filename.endswith(ALLOWED_EXTENSIONS):
            return Response(
                {"error": "Accepted formats are PDF, DOC, and DOCX."},
                status=status.HTTP_400_BAD_REQUEST
            )
        if file.size > MAX_UPLOAD_BYTES:
            return Response(
                {"error": "Maximum file size is 25 MB."},
                status=status.HTTP_400_BAD_REQUEST
            )

        meta = extract_submission_meta(request)
        if not meta['paper_title']:
            return Response({"error": "Paper title is required."}, status=status.HTTP_400_BAD_REQUEST)

        is_b2b_submission = request.data.get('is_b2b') == 'true' or request.data.get('is_b2b') is True

        word_count = get_word_count(file)
        config = PricingConfig.objects.last()
        if not config:
            config = PricingConfig.objects.create(**default_pricing_kwargs())

        package, total_price, package_flags = resolve_package(request.data, config)

        if is_b2b_submission:
            # Check B2B eligibility
            if not user.college:
                return Response({"error": "Your account is not associated with any college B2B credits."}, status=status.HTTP_400_BAD_REQUEST)
            if user.college.credits < 1:
                return Response({"error": "Insufficient credits remaining for your college account. Please contact college admin."}, status=status.HTTP_400_BAD_REQUEST)
            
            # Deduct 1 credit
            college = user.college
            college.credits -= 1
            college.save()

            # Check low credit warning (credits < 20% of allocated)
            if college.allocated_credits > 0 and (college.credits / college.allocated_credits) < 0.20:
                # Log warning
                print(f"[ALERT] College '{college.college_name}' credits are low! ({college.credits} remaining out of {college.allocated_credits})")
                # Send email to college admin if email exists
                if college.contact_email:
                    try:
                        send_mail(
                            'Low Credit Alert - NovelCheckr Platform',
                            f"Dear Admin,\n\nYour college account credits are running low. Remaining balance: {college.credits} out of {college.allocated_credits}.\n\nPlease renew/top up credits from the dashboard.\n\nBest regards,\nNovelCheckr Team",
                            'admin@novelcheckr.com',
                            [college.contact_email],
                            fail_silently=True
                        )
                    except Exception as e:
                        print(f"Failed to send low credit email: {e}")

            # Create Order
            order = Order.objects.create(
                user=user,
                document=file,
                word_count=word_count,
                price=0.00,
                status='Submitted',
                package_tier=package,
                is_express=package_flags['is_express'],
                has_editing_suggestions=package_flags['has_editing_suggestions'],
                is_b2b=True,
                college=college,
                department=user.department or 'General',
                **meta,
            )

            serializer = OrderSerializer(order)
            return Response(serializer.data, status=status.HTTP_201_CREATED)
            
        else:
            # Create Order (status: Pending Payment, waiting for payment completion)
            order = Order.objects.create(
                user=user,
                document=file,
                word_count=word_count,
                price=total_price,
                status='Pending Payment',
                package_tier=package,
                is_express=package_flags['is_express'],
                has_editing_suggestions=package_flags['has_editing_suggestions'],
                is_b2b=False,
                **meta,
            )

            serializer = OrderSerializer(order)
            return Response(serializer.data, status=status.HTTP_201_CREATED)


class OrderDetailView(generics.RetrieveAPIView):
    queryset = Order.objects.select_related('user', 'college', 'payment')
    serializer_class = OrderSerializer


class AddEditingSuggestionsView(APIView):
    def post(self, request, pk):
        try:
            order = Order.objects.get(pk=pk)
        except Order.DoesNotExist:
            return Response({"error": "Order not found"}, status=status.HTTP_404_NOT_FOUND)

        if order.has_editing_suggestions:
            return Response({"message": "Editing suggestions already added"}, status=status.HTTP_200_OK)

        config = PricingConfig.objects.last()
        upsell_fee = config.editing_suggestions_fee if config else default_pricing_kwargs()['editing_suggestions_fee']
        
        order.has_editing_suggestions = True
        order.package_tier = 'complete'
        order.price += upsell_fee
        order.save()

        # If there is a payment record, update its amount as well
        if hasattr(order, 'payment'):
            payment = order.payment
            payment.amount += upsell_fee
            payment.save()

        return Response({
            "message": "Editing suggestions added successfully",
            "price": order.price
        })


class DownloadReportView(APIView):
    permission_classes = [permissions.AllowAny] # Anyone with the secure signed URL can download

    def get(self, request, pk):
        try:
            order = Order.objects.get(pk=pk)
        except Order.DoesNotExist:
            return Response({"error": "Order not found"}, status=status.HTTP_404_NOT_FOUND)

        token = request.query_params.get('token')
        if not token:
            return Response({"error": "Secure download token is missing"}, status=status.HTTP_400_BAD_REQUEST)

        # Verify Signature & Age
        try:
            # Max age of 48 hours (48 * 3600 = 172800 seconds)
            unsigned_id = signer.unsign(token, max_age=172800)
            if int(unsigned_id) != order.id:
                return Response({"error": "Invalid token signature"}, status=status.HTTP_400_BAD_REQUEST)
        except SignatureExpired:
            return Response({"error": "This report link has expired (valid for 48 hours only)"}, status=status.HTTP_400_BAD_REQUEST)
        except BadSignature:
            return Response({"error": "Invalid secure download link signature"}, status=status.HTTP_400_BAD_REQUEST)

        # Check report_uploaded_at as a secondary safety check
        if order.report_uploaded_at:
            delta = timezone.now() - order.report_uploaded_at
            if delta.total_seconds() > 172800:
                return Response({"error": "Report download window of 48 hours has expired"}, status=status.HTTP_400_BAD_REQUEST)

        if not order.report_file:
            return Response({"error": "Report file has not been uploaded yet"}, status=status.HTTP_400_BAD_REQUEST)

        # Serve file response
        file_path = order.report_file.path
        if not os.path.exists(file_path):
            return Response({"error": "Report file not found on server storage"}, status=status.HTTP_404_NOT_FOUND)

        return FileResponse(open(file_path, 'rb'), content_type='application/pdf')


class OrderInvoiceView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request, pk):
        try:
            order = Order.objects.get(pk=pk)
        except Order.DoesNotExist:
            return Response({"error": "Order not found"}, status=status.HTTP_404_NOT_FOUND)

        # Check permissions
        if request.user.role != 'super_admin' and order.user != request.user:
            return Response({"error": "You do not have access to this invoice"}, status=status.HTTP_403_FORBIDDEN)

        # Import reportlab details
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        from io import BytesIO

        try:
            pdf_buffer = BytesIO()
            doc = SimpleDocTemplate(pdf_buffer, pagesize=letter, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
            story = []

            styles = getSampleStyleSheet()
            
            NAVY = colors.HexColor("#0c2340")
            GOLD = colors.HexColor("#c5a572")
            SLATE = colors.HexColor("#1e293b")
            MUTED = colors.HexColor("#64748b")
            LIGHT_BG = colors.HexColor("#f8fafc")
            BORDER_COLOR = colors.HexColor("#e2e8f0")
            GOLD_BG = colors.HexColor("#fefce8")

            # Custom Styles
            brand_title = ParagraphStyle(
                'InvoiceBrandTitle',
                parent=styles['Normal'],
                fontName="Helvetica-Bold",
                fontSize=22,
                textColor=NAVY,
                leading=26
            )
            invoice_title = ParagraphStyle(
                'InvoiceDocTitle',
                parent=styles['Normal'],
                fontName="Helvetica-Bold",
                fontSize=18,
                textColor=NAVY,
                alignment=2,
                leading=22
            )
            body_style = ParagraphStyle(
                'InvoiceBody',
                parent=styles['Normal'],
                fontName="Helvetica",
                fontSize=9,
                textColor=SLATE,
                leading=13
            )
            body_bold = ParagraphStyle(
                'InvoiceBodyBold',
                parent=styles['Normal'],
                fontName="Helvetica-Bold",
                fontSize=9,
                textColor=SLATE,
                leading=13
            )
            table_header_style = ParagraphStyle(
                'InvoiceTableHeader',
                parent=styles['Normal'],
                fontName="Helvetica-Bold",
                fontSize=9.5,
                textColor=colors.white,
                leading=12
            )

            # Metadata preparation
            doc_name = os.path.basename(order.document.name) if order.document else "Document"
            user_full_name = order.user.get_full_name() or order.user.username
            user_phone = getattr(order.user, 'phone', 'N/A') or 'N/A'
            payment_ref = getattr(order, 'payment', None) and (order.payment.razorpay_payment_id or order.payment.transaction_id) or ('B2B Credit Allocation' if order.is_b2b else 'Direct Transaction')

            # 1. Header Row
            header_table_data = [
                [
                    Paragraph("<b>NOVELCHECKR</b><br/><font color='#c5a572'><b>ACADEMIC INTEGRITY & VERIFICATION INTELLIGENCE</b></font>", brand_title),
                    Paragraph("TAX INVOICE<br/><font size=9 color='#64748b'>Invoice No: <b>INV-%06d</b><br/>Date: <b>%s</b></font>" % (order.id, order.created_at.strftime('%d %b %Y')), invoice_title)
                ]
            ]
            header_table = Table(header_table_data, colWidths=[310, 230])
            header_table.setStyle(TableStyle([
                ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                ('LEFTPADDING', (0,0), (-1,-1), 0),
                ('RIGHTPADDING', (0,0), (-1,-1), 0),
            ]))
            story.append(header_table)
            story.append(Spacer(1, 12))

            # 2. Divider Line (Navy accent line)
            story.append(Table([['']], colWidths=[540], rowHeights=[2], style=TableStyle([('BACKGROUND', (0,0), (-1,-1), NAVY)])))
            story.append(Spacer(1, 14))

            # 3. Billing Info Cards (Billed To & Invoice Metadata)
            info_data = [
                [
                    Paragraph(
                        f"<font color='#0c2340'><b>BILLED TO</b></font><br/><br/>"
                        f"Client Name: <b>{user_full_name}</b><br/>"
                        f"Email: {order.user.email}<br/>"
                        f"Phone: {user_phone}<br/>"
                        f"Department: {order.department or 'General'}",
                        body_style
                    ),
                    Paragraph(
                        f"<font color='#0c2340'><b>ISSUER & PAYMENT INFO</b></font><br/><br/>"
                        f"Issued By: <b>NovelCheckr Platform Services</b><br/>"
                        f"Billing Model: <b>{'B2B Institutional Credit' if order.is_b2b else 'B2C Online Payment'}</b><br/>"
                        f"Payment Ref: <b>{payment_ref}</b><br/>"
                        f"Order Status: <b>{order.status}</b>",
                        body_style
                    )
                ]
            ]
            info_table = Table(info_data, colWidths=[264, 264])
            info_table.setStyle(TableStyle([
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('BACKGROUND', (0,0), (-1,-1), LIGHT_BG),
                ('PADDING', (0,0), (-1,-1), 12),
                ('BOX', (0,0), (0,0), 1, BORDER_COLOR),
                ('BOX', (1,0), (1,0), 1, BORDER_COLOR),
            ]))
            story.append(info_table)
            story.append(Spacer(1, 14))

            # 4. Manuscript Details Banner
            paper_title_display = getattr(order, 'paper_title', '') or doc_name
            paper_meta_data = [
                [
                    Paragraph(
                        f"<font color='#0c2340'><b>VERIFICATION SUBJECT & MANUSCRIPT META</b></font><br/>"
                        f"<font color='#64748b'>Title:</font> <b>{paper_title_display}</b> &nbsp;&nbsp;•&nbsp;&nbsp; "
                        f"<font color='#64748b'>File:</font> <b>{doc_name}</b> &nbsp;&nbsp;•&nbsp;&nbsp; "
                        f"<font color='#64748b'>Word Count:</font> <b>{order.word_count or 'N/A'} words</b>",
                        body_style
                    )
                ]
            ]
            paper_meta_table = Table(paper_meta_data, colWidths=[540])
            paper_meta_table.setStyle(TableStyle([
                ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                ('BACKGROUND', (0,0), (-1,-1), GOLD_BG),
                ('PADDING', (0,0), (-1,-1), 10),
                ('BOX', (0,0), (-1,-1), 1, colors.HexColor("#fef08a")),
            ]))
            story.append(paper_meta_table)
            story.append(Spacer(1, 16))

            # 5. Line Items Table with GST Breakdown
            package = order.package_tier or ('complete' if order.has_editing_suggestions else 'improve' if order.is_express else 'check')
            package_name = PACKAGE_LABELS.get(package, 'Check — Similarity Check')

            if order.is_b2b:
                table_data = [
                    [Paragraph("Service Description", table_header_style), Paragraph("Taxable Value", table_header_style), Paragraph("GST (18%)", table_header_style), Paragraph("Total Billed", table_header_style)],
                    [
                        Paragraph(f"<b>{package_name}</b><br/><font size=8 color='#64748b'>Institutional B2B Verification Check</font>", body_style),
                        Paragraph("—", body_style),
                        Paragraph("—", body_style),
                        Paragraph("<b>1 B2B Credit</b>", body_style),
                    ],
                    [
                        Paragraph("<b>Total Amount Billed</b>", body_bold),
                        Paragraph("", body_style),
                        Paragraph("", body_style),
                        Paragraph("<b>1 B2B Credit</b>", body_bold),
                    ],
                ]
            else:
                taxable, gst, gross = gst_breakdown(order.price)
                table_data = [
                    [Paragraph("Service Description", table_header_style), Paragraph("Taxable Value", table_header_style), Paragraph("GST (18%)", table_header_style), Paragraph("Total Paid (INR)", table_header_style)],
                    [
                        Paragraph(f"<b>{package_name}</b><br/><font size=8 color='#64748b'>Includes GST & Report Generation</font>", body_style),
                        Paragraph("₹ %.2f" % taxable, body_style),
                        Paragraph("₹ %.2f" % gst, body_style),
                        Paragraph("₹ %.2f" % gross, body_style),
                    ],
                    [
                        Paragraph("<b>Total Amount Paid (GST Included)</b>", body_bold),
                        Paragraph("₹ %.2f" % taxable, body_bold),
                        Paragraph("₹ %.2f" % gst, body_bold),
                        Paragraph("<b>₹ %.2f</b>" % gross, body_bold),
                    ],
                ]

            summary_table = Table(table_data, colWidths=[240, 100, 100, 100])
            summary_table.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,0), NAVY),
                ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                ('BOTTOMPADDING', (0,0), (-1,-1), 8),
                ('TOPPADDING', (0,0), (-1,-1), 8),
                ('GRID', (0,0), (-1,-2), 0.5, BORDER_COLOR),
                ('LINEABOVE', (0,-1), (-1,-1), 1.5, NAVY),
                ('BACKGROUND', (0,-1), (-1,-1), LIGHT_BG),
            ]))
            story.append(summary_table)
            story.append(Spacer(1, 24))

            # 6. Computer Generated Seal & Footer
            footer_data = [
                [
                    Paragraph(
                        "<font color='#64748b' size=8><b>Official Tax Receipt & Verification Audit:</b><br/>"
                        "This is an official computer-generated digital tax invoice issued by NovelCheckr. "
                        "No physical signature is required. For support, contact <u>support@novelcheckr.com</u>.</font>",
                        body_style
                    ),
                    Paragraph(
                        "<font color='#0c2340' size=8.5><b>NOVELCHECKR</b></font><br/>"
                        "<font color='#c5a572' size=7.5><b>VERIFIED AUDIT RECEIPT</b></font>",
                        ParagraphStyle('SealText', parent=styles['Normal'], alignment=1)
                    )
                ]
            ]
            footer_table = Table(footer_data, colWidths=[410, 130])
            footer_table.setStyle(TableStyle([
                ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                ('BACKGROUND', (1,0), (1,0), LIGHT_BG),
                ('BOX', (1,0), (1,0), 1, GOLD),
                ('PADDING', (1,0), (1,0), 8),
            ]))
            story.append(footer_table)

            doc.build(story)
            pdf_buffer.seek(0)
            pdf_bytes = pdf_buffer.getvalue()
            pdf_buffer.close()

            response = HttpResponse(pdf_bytes, content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="Invoice_{order.id}.pdf"'
            return response
            
        except Exception as e:
            print(f"Error generating invoice for order {pk}: {str(e)}")
            import traceback
            traceback.print_exc()
            return Response({"error": f"Failed to generate invoice: {str(e)}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class SuperAdminOrderQueueView(generics.ListAPIView):
    serializer_class = OrderSerializer
    permission_classes = [IsSuperAdmin]

    def get_queryset(self):
        # Return all pending orders (exclude Report Ready and Pending Payment), sorted express priority first
        return Order.objects.exclude(status__in=['Report Ready', 'Pending Payment']).select_related('user', 'college', 'payment').order_by('-is_express', '-created_at')


class SuperAdminUpdateOrderView(APIView):
    permission_classes = [IsSuperAdmin]

    def post(self, request, pk):
        try:
            order = Order.objects.get(pk=pk)
        except Order.DoesNotExist:
            return Response({"error": "Order not found"}, status=status.HTTP_404_NOT_FOUND)

        action = request.data.get('action') # 'start_processing' or 'complete'
        
        if action == 'start_processing':
            order.status = 'Processing'
            order.save()
            try:
                from services.firestore_service import save_order_to_firestore
                save_order_to_firestore(order)
            except Exception as e:
                print(f"Firestore Order Sync Error: {e}")
            return Response({
                "message": "Order marked as Processing",
                "status": order.status
            })
            
        elif action == 'complete':
            similarity_score = request.data.get('similarity_score')
            report_file = request.FILES.get('report_file')
            
            if similarity_score is None:
                return Response({"error": "Similarity score is required"}, status=status.HTTP_400_BAD_REQUEST)
            if not report_file:
                return Response({"error": "Report PDF file is required"}, status=status.HTTP_400_BAD_REQUEST)

            try:
                similarity_score = float(similarity_score)
            except ValueError:
                return Response({"error": "Similarity score must be a number"}, status=status.HTTP_400_BAD_REQUEST)

            order.similarity_score = similarity_score
            order.report_file = report_file
            order.status = 'Report Ready'
            order.report_uploaded_at = timezone.now()
            order.save()

            try:
                from services.firestore_service import save_order_to_firestore
                save_order_to_firestore(order)
            except Exception as e:
                print(f"Firestore Order Sync Error: {e}")

            # Generate expiring signed download URL
            token = signer.sign(str(order.id))
            # Create absolute URL or relative download link
            secure_link = f"/api/orders/{order.id}/download-report/?token={token}"

            # Simulate Notifications (Email & WhatsApp)
            print(f"[NOTIFICATION] Email and WhatsApp sent to user '{order.user.username}' (phone: {order.user.phone or 'N/A'}, email: {order.user.email})")
            print(f"[WHATSAPP SIMULATION] Message: Your document integrity report for {os.path.basename(order.document.name)} is ready! Similarity: {similarity_score}%. Download link valid for 48 hours: {secure_link}")

            try:
                # Send Email
                send_mail(
                    'Your Integrity Verification Report is Ready!',
                    f"Hi {order.user.username},\n\nYour document '{os.path.basename(order.document.name)}' has been verified.\nSimilarity Score: {similarity_score}%\n\nYou can access the secure download link below. Note: This link is valid only for 48 hours:\n{request.build_absolute_uri(secure_link)}\n\nBest regards,\nNovelCheckr Support Team",
                    'support@novelcheckr.com',
                    [order.user.email],
                    fail_silently=True
                )
            except Exception as ex:
                print(f"Failed to send notification email: {ex}")

            return Response({
                "message": "Order marked as Complete. Notifications triggered.",
                "status": order.status,
                "similarity_score": order.similarity_score,
                "secure_download_url": secure_link
            })
            
        else:
            return Response({"error": "Invalid action. Must be 'start_processing' or 'complete'"}, status=status.HTTP_400_BAD_REQUEST)


class PricingConfigView(APIView):
    # Retrieve configuration (available to all logged in users)
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        config = PricingConfig.objects.last()
        if not config:
            config = PricingConfig.objects.create(**default_pricing_kwargs())
        serializer = PricingConfigSerializer(config)
        payload = serializer.data
        payload['gst_rate'] = 18
        payload['packages'] = package_catalog(config)
        # Always expose GST-inclusive public prices, never leftover per-word 0.50
        payload['per_word_rate'] = payload['packages']['check']['price']
        payload['express_fee'] = payload['packages']['improve']['price']
        payload['editing_suggestions_fee'] = payload['packages']['complete']['price']
        return Response(payload)

    def post(self, request):
        # Update configuration (superadmin only)
        if request.user.role != 'super_admin':
            return Response({"error": "Only super admins can modify pricing configuration"}, status=status.HTTP_403_FORBIDDEN)
            
        config = PricingConfig.objects.last()
        if not config:
            config = PricingConfig()
            
        serializer = PricingConfigSerializer(config, data=request.data, partial=True)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
